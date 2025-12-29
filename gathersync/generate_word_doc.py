#!/usr/bin/env python3
"""Generate GatherSync instructions as Word document with embedded QR code"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Helvetica'
font.size = Pt(11)

# Title
title = doc.add_heading('Welcome to GatherSync! 🎉', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Your Personal Event Coordination Assistant')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# Introduction
intro = doc.add_paragraph(
    'Peter has invited you to test GatherSync - a new app designed to make coordinating '
    'group events easier. This app helps you find the perfect date when everyone is available.'
)

doc.add_paragraph()

# Quick Start Guide
doc.add_heading('Quick Start Guide', level=2)

# Step 1
doc.add_heading('Step 1: Install Expo Go', level=3)
doc.add_paragraph('iPhone Users:', style='List Bullet')
doc.add_paragraph('Open the App Store', style='List Number')
doc.add_paragraph('Search for "Expo Go"', style='List Number')
doc.add_paragraph('Install the app (it\'s free)', style='List Number')

doc.add_paragraph()

doc.add_paragraph('Android Users:', style='List Bullet')
doc.add_paragraph('Open the Google Play Store', style='List Number')
doc.add_paragraph('Search for "Expo Go"', style='List Number')
doc.add_paragraph('Install the app (it\'s free)', style='List Number')

doc.add_paragraph()

# Step 2
doc.add_heading('Step 2: Open GatherSync', level=3)
doc.add_paragraph('Option A: Scan the QR Code (Easiest!)', style='Heading 4')
doc.add_paragraph('Open the Expo Go app', style='List Number')
doc.add_paragraph('Tap "Scan QR code"', style='List Number')
doc.add_paragraph('Point your camera at the QR code below:', style='List Number')

doc.add_paragraph()

# Add QR code image
try:
    qr_paragraph = doc.add_paragraph()
    qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = qr_paragraph.add_run()
    run.add_picture('gathersync-qr.png', width=Inches(3))
except Exception as e:
    doc.add_paragraph(f'[QR Code image will be inserted here: {e}]')

doc.add_paragraph()
doc.add_page_break()

# Option B
doc.add_paragraph('Option B: Enter URL Manually', style='Heading 4')
doc.add_paragraph('Open the Expo Go app', style='List Number')
doc.add_paragraph('Tap "Enter URL manually" at the bottom', style='List Number')
doc.add_paragraph('Copy and paste this URL:', style='List Number')

url_para = doc.add_paragraph()
url_run = url_para.add_run('https://8081-ienb1rj930k0x92csc3x6-a41ba8ee.manus-asia.computer')
url_run.font.name = 'Courier New'
url_run.font.size = Pt(9)
url_run.font.color.rgb = RGBColor(37, 99, 235)

doc.add_paragraph('Tap "Connect"', style='List Number')

doc.add_paragraph()

# Step 3
doc.add_heading('Step 3: Log In to Enable Cloud Sync', level=3)
doc.add_paragraph('When GatherSync opens, you\'ll see a "Cloud Sync Available" banner', style='List Number')
doc.add_paragraph('Tap the "Log In" button', style='List Number')
doc.add_paragraph('Create an account or use OAuth (Google/GitHub)', style='List Number')
doc.add_paragraph('After logging in, you\'ll be redirected back to the app', style='List Number')

doc.add_paragraph()

# Step 4
doc.add_heading('Step 4: Sync Your Events', level=3)
doc.add_paragraph('You should now see "Cloud sync enabled" banner (green)', style='List Number')
doc.add_paragraph('Tap the "Sync Now" button', style='List Number')
doc.add_paragraph('Wait 10-20 seconds while syncing', style='List Number')
doc.add_paragraph('You should see 2 events appear:', style='List Number')

events_list = doc.add_paragraph(style='List Bullet')
events_list.add_run('AI Guys').bold = True
events_list.add_run(' - Monthly flexible event')

events_list2 = doc.add_paragraph(style='List Bullet')
events_list2.add_run('Guru Breakfast').bold = True
events_list2.add_run(' - Monthly event')

doc.add_paragraph()
doc.add_page_break()

# Step 5
doc.add_heading('Step 5: Mark Your Availability', level=3)

doc.add_paragraph('For Flexible Events (AI Guys):', style='Heading 4')
doc.add_paragraph('Tap on the AI Guys event', style='List Number')
doc.add_paragraph('Find your name in the participants list', style='List Number')
doc.add_paragraph('Tap on your name', style='List Number')
doc.add_paragraph('Mark the dates you\'re available or unavailable', style='List Number')
doc.add_paragraph('The app will help find the best date when most people can attend', style='List Number')

doc.add_paragraph()

doc.add_paragraph('For Fixed Events (Guru Breakfast):', style='Heading 4')
doc.add_paragraph('Tap on the Guru Breakfast event', style='List Number')
doc.add_paragraph('Find your name in the participants list', style='List Number')
doc.add_paragraph('Tap on your name', style='List Number')
doc.add_paragraph('Select your RSVP status:', style='List Number')

doc.add_paragraph('✅ Attending', style='List Bullet')
doc.add_paragraph('❌ Not Attending', style='List Bullet')
doc.add_paragraph('❓ No Response (default)', style='List Bullet')

doc.add_paragraph()

# Step 6
doc.add_heading('Step 6: Sync Your Changes', level=3)
doc.add_paragraph('After marking your availability or RSVP:')
doc.add_paragraph('Go back to the Events list (tap the back arrow)', style='List Number')
doc.add_paragraph('Tap "Sync Now" again', style='List Number')
doc.add_paragraph('Your responses will be uploaded to the cloud', style='List Number')
doc.add_paragraph('Everyone else will see your updates when they sync', style='List Number')

doc.add_paragraph()
doc.add_page_break()

# Important Tips
doc.add_heading('Important Tips', level=2)
tip1 = doc.add_paragraph('✅ ', style='List Bullet')
tip1.add_run('Always sync after making changes').bold = True
tip1.add_run(' - Your availability/RSVP won\'t be shared until you tap "Sync Now"')

tip2 = doc.add_paragraph('✅ ', style='List Bullet')
tip2.add_run('Sync regularly to see updates').bold = True
tip2.add_run(' - Other people\'s responses will appear when you sync')

tip3 = doc.add_paragraph('✅ ', style='List Bullet')
tip3.add_run('Keep Expo Go installed').bold = True
tip3.add_run(' - You\'ll need it to access GatherSync until we release the standalone app')

tip4 = doc.add_paragraph('✅ ', style='List Bullet')
tip4.add_run('Internet connection required').bold = True
tip4.add_run(' - Syncing requires an active internet connection')

doc.add_paragraph()

# Need Help
doc.add_heading('Need Help?', level=2)
doc.add_paragraph('If you encounter any issues:')
doc.add_paragraph('Force close and reopen - Swipe up from app switcher and reopen Expo Go', style='List Number')
doc.add_paragraph('Check your internet connection - Sync requires WiFi or mobile data', style='List Number')
doc.add_paragraph('Contact Peter - He\'s testing this app and wants your feedback!', style='List Number')

doc.add_paragraph()

# What to Test
doc.add_heading('What to Test', level=2)
doc.add_paragraph('Peter would love your feedback on:')
doc.add_paragraph('✅ Does the app load correctly?', style='List Bullet')
doc.add_paragraph('✅ Can you see the events after syncing?', style='List Bullet')
doc.add_paragraph('✅ Can you mark your availability/RSVP?', style='List Bullet')
doc.add_paragraph('✅ Do your changes sync back to other users?', style='List Bullet')
doc.add_paragraph('✅ Is the app easy to use?', style='List Bullet')
doc.add_paragraph('✅ What features would make it more useful?', style='List Bullet')

doc.add_paragraph()

# Privacy
doc.add_heading('Privacy & Data', level=2)
doc.add_paragraph('Your data is stored securely in the cloud', style='List Bullet')
doc.add_paragraph('Only members of your events can see your responses', style='List Bullet')
doc.add_paragraph('You can delete your account and data at any time', style='List Bullet')
doc.add_paragraph('This is a test version - please don\'t enter sensitive information', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()

# Thank you
thank_you = doc.add_paragraph('Thank you for testing GatherSync!')
thank_you.alignment = WD_ALIGN_PARAGRAPH.CENTER
thank_you.runs[0].bold = True
thank_you.runs[0].font.size = Pt(14)

feedback = doc.add_paragraph('Your feedback will help make this app better for everyone.')
feedback.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

contact = doc.add_paragraph('Questions? Contact Peter Scarfo')
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.runs[0].italic = True

doc.add_paragraph()

copyright_para = doc.add_paragraph('© 2025 Peter Scarfo. All rights reserved.')
copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
copyright_para.runs[0].font.size = Pt(9)

# Save document
doc.save('GatherSync-Instructions.docx')
print('✅ Word document generated successfully: GatherSync-Instructions.docx')
